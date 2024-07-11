import docx
temp = docx.Document("./Word Files/BEER SYSTEM_Perlick_4414W.docx")
doc = docx.Document()
fullText = []
            
p = doc.add_paragraph('')

print(len(temp.paragraphs))
for para in temp.paragraphs:
    print(para.text)
    p_runs = []
    addRuns = False
    beginning = True
    for runS in para.runs:
        if runS.font.color.rgb != docx.shared.RGBColor(0x00, 0x00, 0x00) and runS.font.color.rgb != None:
            print(runS.font.color.rgb)
            if fullText:
                p.add_run('\n'.join(fullText) + '\n')
                print(fullText)
                fullText = []
            if beginning:
                beginning = False
                #p.add_run('\n')
            if p_runs:
                print(p_runs)
                p.add_run(''.join(p_runs))
            
            p_runs = []
            addRuns = True
            redRun = p.add_run(runS.text)
            redRun.font.color.rgb = runS.font.color.rgb
        else:
            print("Everything's normal")
            p_runs.append(runS.text)
    if addRuns:
        p.add_run(''.join(p_runs))
        #p.add_run('\n')
        addRuns = False
    else:
        fullText.append(para.text)       
    if not fullText:
        p.add_run('\n')
    p_runs = []
    
p.add_run('\n'.join(fullText))
doc.save("TEST.docx")