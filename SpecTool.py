#SDI Specs Formatting Tool

#Imports
import os
from os import listdir
from os.path import isfile, join
import sqlite3
import re
import sys
import traceback
import time 
import tkinter as tk
from tkinter import filedialog,ttk


from fast_edit_distance import edit_distance
import docx as d
import openpyxl as opx
import num2words as n2m
from PIL import Image, ImageTk

import LogErrors
import FindHeaders
import SpecDB as db


def resource_path(rel_path):
    try:
        base_path = sys.MEIPASS 
    except Exception:
        base_path = os.path.abspath('./_internal')
    return os.path.join(base_path, rel_path)


#input/output locations
inputFilepath = ""
outputFilepath = ""
excelFilepath = ""

#tkinter root window
root = tk.Tk()

root.title("SDI Specs Formatting Tool")
root.geometry("800x500")

menubar = tk.Menu()
menubar.add_command(label = "Help", command = lambda:os.startfile(resource_path('Help.html')))
root.config(menu=menubar)

#Catch unhandled errors and log them
def handle_exception(exc,val,tb):
    top=tk.Toplevel(root)
    top.geometry("800x400")
    top.title("Error")
    tk.Label(top, text="ERROR: \n", font=(25)).pack()
    tk.Label(top, text="".join(traceback.format_exception(exc,val,tb))).pack()
    tk.Button(top, text="OK", command=top.destroy).pack()
    root.wait_window(top)
    LogErrors.handle_exception(exc,val,tb)
root.report_callback_exception = handle_exception

ico = Image.open(resource_path("SDI Logo.jpg"))


photo = ImageTk.PhotoImage(ico)
root.wm_iconphoto(False, photo)

#Copy Specs Using Filepath as a Key for the Specs DB
def copySpecs(tempDocPath, p, highlight, cur):
    specText = None
    try:
        specText = cur.execute("SELECT text FROM spec WHERE doc='" + tempDocPath + "'").fetchone()[0].split("\n")
    except:
        print("Specs not found for: " + tempDocPath)
        return -1
    #Begin copying after header
    i = 0
    while i<len(specText) and "Utilities" not in specText[i]:
        i = i + 1
        
    #Divide sections by font color
    if (len(("\n".join(specText[i+1:])).split('~')) - 1)%3 != 0:
        specText = ["\n".join(specText[i+1:])]
    else:
        specText = ("\n".join(specText[i+1:])).split('~')

    j=0
    #Copy all text, if font is colored, use that font color
    while j < len(specText):
        #Uncolored Section
        if j%3 == 0:
            if(highlight):
                p.add_run(specText[j]).font.highlight_color = d.enum.text.WD_COLOR_INDEX.YELLOW
            else:
                p.add_run(specText[j])
        #Colored Section 
        else:
            colorRun = p.add_run(specText[j])
            try:
                #Parse Font Color into RGB color components            
                r=int(specText[j+1][:2], 16)
                g=int(specText[j+1][2:4],16)
                b=int(specText[j+1][4:6],16)
                colorRun.font.color.rgb = d.shared.RGBColor(r,g,b)
                j = j+1
            except Exception:
                pass
            
            if highlight:
                colorRun.font.highlight_color = d.enum.text.WD_COLOR_INDEX.YELLOW
        j=j+1
    return None

#Get selected item from treeview and copy filepath (into chosenSpec) then destroy popup window
def ChooseSpec(chosenSpec, popup, tv):
    item = tv.item(tv.focus())
    if len(item['values']) < 2:
        print("Make a Selection")
    else:
        chosenSpec.set(item['values'][3])
        popup.destroy()
        

#Find matching Specs for each item, and enter into .xlsx file highlighting partial matches and non-matches
def findSpecs(msgLabel):
    con = None
    dbPath = resource_path('specsDB.db')
    try:
        con=sqlite3.connect("file:"+dbPath+"?mode=rw", uri=True)
    except sqlite3.OperationalError:
        con=sqlite3.connect(dbPath)
        cur = con.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS item (desc, manu, model, doc)")
        cur.execute("CREATE TABLE IF NOT EXISTS spec (doc, text, modTime)")
        con.commit()
    cur = con.cursor()

    #Make sure input/output paths exist
    try:
        wb = opx.load_workbook(inputFilepath, read_only=True)
    except Exception as e:
        print("Input File Not Found... Please Check Input Filepath")
        msgLabel.config(text="Error: Input File not found")
        return
    if outputFilepath == "":
        msgLabel.config(text="Error: Select Output Location")
        return

    wbNew = opx.Workbook()
    newSheet = wbNew.active
    sheet = wb.active
    
    yellowFill = opx.styles.PatternFill(start_color = 'FFFF00', end_color = 'FFFF00', fill_type = 'solid')
    redFill = opx.styles.PatternFill(start_color = 'FF0000', end_color = 'FF0000', fill_type = 'solid')
    noFill = opx.styles.PatternFill(start_color = 'FFFFFF', end_color = 'FFFFFF', fill_type = 'solid')

    hDict = FindHeaders.FindHeaders(sheet)
    headers = ['qty','description','remarks','manuf.','model']

    missing = []
    for header in headers:
        if header not in hDict:
            missing.append(header)
    if missing:
        msgLabel.config("The following header(s) may be missing: " + ', '.join(missing))
        return
    
    for row in sheet.rows:
        
        #Skip if location header, spare number, existing item, or not in contract
        if row[hDict['qty']].value == None or "spare" in str(row[hDict['description']].value).lower() or (row[hDict['remarks']].value != None and ("by" in str(row[hDict['remarks']].value).lower() or "exist" in str(row[hDict['remarks']].value).lower())):
            continue

        #Collect Description, Manufacturer, and Model No. for matching a Spec ".docx" file
        else:
            ambiguousModels = ["custom", "custom design"] #Non-Unique Model No's
            specData = []
            if row[hDict['remarks']].value != None and "CUSTOM FABRICATION" in str(row[hDict['remarks']].value).upper():
                specData = [row[hDict['description']].value, "Custom Fabrication", ""]
            else:
                specData = [row[hDict['description']].value, row[hDict['manuf.']].value, str(row[hDict['model']].value).replace('/', '-').replace('|','-')]

            #Add row to Excel file with Description, Manufacturer, Model No., and (space for) '.docx' filepath/hyperlink
            newSheet.append([specData[0], specData[1], specData[2], ''])
            rowIndex = newSheet.max_row #current row
            
            matches = []
            #Search DB for a 'exact' match
            if specData[1] == "Custom Fabrication":
                matches = cur.execute("SELECT doc FROM item WHERE desc='" + str(specData[0]).replace("'","''").replace('"','""') + "' COLLATE NOCASE AND manu = 'Custom Fabrication' COLLATE NOCASE").fetchall()
            else:
                matches = cur.execute("SELECT doc FROM item WHERE model='" + str(specData[2]).replace("'","''").replace('"','""') + "'").fetchall()

            #If there was a match, add a link to the xlsx file
            if matches:                
                newSheet[rowIndex][3].value = "=HYPERLINK(\"[" + matches[0][0] + "]\",\""+ matches[0][0].split('/')[len(matches[0][0].split('/'))-1].split('.docx')[0] +"\")"
                for i in range(0,4):
                    newSheet[rowIndex][i].fill = noFill    

            #Else look for partial matches
            else:    
                #Check for partially matching model and manufacturer
                if (specData[1] != "Custom Fabrication" and str(specData[2]).lower() not in ambiguousModels):
                    matches = cur.execute("SELECT desc, model, doc FROM item WHERE model LIKE '%" + str(specData[2]).replace("'","''").replace('"','""') + "%' AND manu LIKE '%" + str(specData[1]).replace("'","''").replace('"','""') + "%'").fetchall()
                #If there still aren't matches search for partially matching description and manufacturer
                if not matches: 
                    matches = cur.execute("SELECT desc, model, doc FROM item WHERE desc LIKE '%" + str(specData[0]).replace("'","''").replace('"','""')+ "%' AND manu LIKE '%"+ str(specData[1]).replace("'","''").replace('"','""') +"%'").fetchall()
                #If there was a match, add a link and highlight yellow
                if matches:

                    def closestMatch(s1, sList, i):
                        minDist = max(len(sList[0][i]), len(s1))
                        closest = sList[0]
                        for s in sList:
                            temp = edit_distance(s1, s[i])
                            if temp < minDist:
                                minDist = temp
                                closest = s
                        return closest
                    
                    if len(matches) > 1:
                        bestMatch = matches[0]
                        if specData[2] == 'Stainless Steel':
                            bestMatch = closestMatch(specData[0], matches, 0)
                        else:
                            bestMatch = closestMatch(specData[2], matches, 1)
                        newSheet[rowIndex][3].value = "=HYPERLINK(\"[" + bestMatch[2] + "]\",\""+bestMatch[2].split('/')[len(bestMatch[2].split('/'))-1].split('.docx')[0] +"\")"
                    else:
                        newSheet[rowIndex][3].value = "=HYPERLINK(\"[" + matches[0][2] + "]\",\""+ matches[0][2].split('/')[len(matches[0][2].split('/'))-1].split('.docx')[0] +"\")"
                    for i in range(0,4):
                        newSheet[rowIndex][i].fill = yellowFill
                #Otherwise highlight red
                else:
                    for i in range(0,4):
                        newSheet[rowIndex][i].fill = redFill
    wb.close()
    #Save new Specs Worksheet
    try:
        wbNew.save(outputFilepath+"\\SpecRefSheet.xlsx")
        global excelFilepath
        excelFilepath = outputFilepath+"/SpecRefSheet.xlsx"
    except PermissionError as e:
        print(e)
        msgLabel.config(text="Error: Cannot Save SpecRefSheet.xlsx while file is open")
        return
    msgLabel.config(text="Successfully Created Spec Ref Sheet")
        
def writeSpecs(msgLabel):
    start_time = time.time()
    
    global excelFilepath
    
    metric = False
    
    specDict = {}

    broken = [] #Holds links from SpecRefSheet that couldn't be found in the database
    
    #Create doc and style
    doc = d.Document(resource_path("Style_Template.docx"))
    
    doc_styles = doc.styles
    
    msgLabel.config(text="Working...")

    #Open Revit output
    try:
        wb = opx.load_workbook(inputFilepath, read_only=True)
    except:
        print("Input File Not Found... Please Check Input Filepath")
        msgLabel.config(text="Error: Input File not found")
        return
    if outputFilepath == '':
        msgLabel.config(text="Error: Select Output location")
        return
    sheet = wb.active

    #Open Specs reference file (optional)
    try:
        wbr = opx.load_workbook(excelFilepath, read_only=True)
    except:
        print("The Excel File was not found... But that's OK")
        excelFilepath = ""
    yellowFill = opx.styles.PatternFill(start_color = 'FFFF00', end_color = 'FFFF00', fill_type = 'solid')
    specRefs = [[],[],[],[],[]]  #Holds [Desc.[], Manufacturer[], Model#[], refFile[], exactMatch?[]] from xl spec ref file

    #If there's a ref sheet, copy item information for later searching
    if excelFilepath:
        refSheet = wbr.active
        for row in refSheet.rows:
            specRefs[0].append(str(row[0].value).lower())
            specRefs[1].append(str(row[1].value).lower())
            specRefs[2].append(str(row[2].value).lower())
            
            if row[3].value == None:
                specRefs[3].append("")
                specRefs[4].append(False)
            else:
                try:
                    if "HYPERLINK" in str(row[3].value): 
                        specRefs[3].append(str(row[3].value).split('\"')[1].replace('[','').replace(']',''))
                    else:
                        path = str(row[3].value).replace('\\','/')
                        if '.docx' not in path: path = path + ".docx"
                        specRefs[3].append(path)
                except:
                    specRefs[3].append("Broken Path :(")
                if row[0].fill == yellowFill:
                    specRefs[4].append(False)
                else:
                    specRefs[4].append(True)

    hDict = FindHeaders.FindHeaders(sheet)
    missing = [] #Holds all headers which were not found

    headers = ['description','qty','equipment','remarks','manuf.','model']
    for head in headers:
        if head not in hDict:
            missing.append(head)
    if missing:
        msgLabel.config(text="The following header(s) are missing: " + ", ".join(missing))
        return

    optionalHeaders = ['hgt._','cw', 'd.w.', 'connection load','voltage', 'phase', 'comments__', 'cfm', 'cfm_', 'hw', 'waste', 'size_', 'in size', 'out size', 'btu\'s','w.c.']
    optionalMetricHeaders = ['hgt. (mm)_', 'cw (mm)','d.w. (mm)', 'connection load', 'voltage', 'phase', 'comments__','m^3/h','m^3/h_','hw (mm)','waste','size (mm)_','in (mm)', 'out (mm)', 'kw', 'mbar']

    missing = [head for head in optionalHeaders if head not in hDict]
    metricMissing = [head for head in optionalMetricHeaders if head not in hDict]

    if len(missing) < len(metricMissing):
        metric = False
        if missing:
            msgLabel.config(text="Warning: The following header(s) may be missing: " + ", ".join(missing))
    else:
        metric = True
        if metricMissing:
            msgLabel.config(text="Warning: The following header(s) may be missing: " + ", ".join(missing))
    
    #Iterate every row in Revit output sheet
    for row in sheet.rows:
    
        #Skip header rows
        if row[hDict['equipment']].value == None or row[hDict['equipment']].value == 'NO' or row[hDict['equipment']].value == 'EQUIPMENT':
            continue
        
        #Add section header (location/area)
        elif row[hDict['qty']].value == None:
            p = doc.add_paragraph('', style = 'Spec_Header')
            p.alignment = 1
            p.add_run(str(row[0].value) + "\n").bold = True

#HEADERS        
        #Create spec header with info from Revit output
        else:
            p = doc.add_paragraph('', style = 'Spec_Header')
            p.alignment = 0

            #Doc formatting
            tab_stops = p.paragraph_format.tab_stops
            tab_stop = tab_stops.add_tab_stop(d.shared.Inches(1.31), d.enum.text.WD_TAB_ALIGNMENT.LEFT)
            tab_stop = tab_stops.add_tab_stop(d.shared.Inches(1.69))
        

            run = "" #To hold all text for each header
            
            #Item Number and Description
            run = run + ("ITEM #" + str(row[hDict['equipment']].value) + ":")
            run = run + ("\t" + str(row[hDict['description']].value))

            #Maybe catches incorrect file related errors
            try:
                row[hDict['remarks']].value != None
                "SPARE NUMBER" in str(row[hDict['remarks']].value)
            except:
                msgLabel.config(text="Error: Specs Not Found. Please check input file is correct.")
                return

            #Check if not in contract
            if row[hDict['remarks']].value != None and ("by vendor" in str(row[hDict['remarks']].value).lower() or "by os&e" in str(row[hDict['remarks']].value).lower() or "by general contractor" in str(row[hDict['remarks']].value).lower() or "by owner" in str(row[hDict['remarks']].value).lower() or "by mep" in str(row[hDict['remarks']].value).lower()):
                run = run + " (NOT IN CONTRACT)"
        
            #Check if existing equipment
            if row[hDict['remarks']].value != None and "EXIST" in str(row[hDict['remarks']].value):
                run = run + " (EXISTING EQUIPMENT)"
            
            #Highlight if shelving unit    
            if "shelv" in str(row[hDict['description']].value).lower():
                r = p.add_run(run)
                run = ""
                r.font.highlight_color = d.enum.text.WD_COLOR_INDEX.YELLOW
        
            #Skip rest if Spare Number
            if "SPARE NUMBER" in str(row[hDict['description']].value):
                p.add_run(run)
                continue

            #Quantity
            run = run + ("\nQuantity:\t")
            try:
                run = run + (n2m.num2words(row[hDict['qty']].value).capitalize() + " (" + str(row[hDict['qty']].value) + ")")
            except:
                run = run + str(row[hDict['qty']].value)
                
            #If not in contract, add pert. data
            if type(row[hDict['remarks']].value) == str and ("by mep" in row[hDict['remarks']].value.lower() or "by vendor" in row[hDict['remarks']].value.lower() or "by os&e" in row[hDict['remarks']].value.lower() or "by general contractor" in row[hDict['remarks']].value.lower() or "by owner" in row[hDict['remarks']].value.lower()):
                run = run+ ("\nPertinent Data:\t")
                if(row[hDict['remarks']].value == None):
                    run = run + add_run("---")
                else:
                    run = run + (str(row[hDict['remarks']].value))
                p.add_run(run)
                continue
        
            #Manufacturer
            run = run + ("\nManufacturer:\t")

            customFab = False
            
            #Check if custom fab
            if(type(row[hDict['remarks']].value) == str and "CUSTOM FABRICATION" in str(row[hDict['remarks']].value).upper()):
                run = run + "Custom Fabrication"
                customFab=True
            else:
                if(row[hDict['manuf.']].value == None):
                    run = run + ("---")
                else:
                    run = run + (str(row[hDict['manuf.']].value))
        
            #Model Number
            run = run + ("\nModel No.:\t")
            if(row[hDict['model']].value == None):
                run = run + ("---")
            else: 
                run = run + (str(row[hDict['model']].value))


            #Pertinent Data
            run = run + ("\nPertinent Data:\t")
        
            #Remove "Custom Fabrication" from pert. data
            if(type(row[hDict['remarks']].value) == str and "CUSTOM FABRICATION" in row[hDict['remarks']].value.upper()):
                temp = "".join(re.split("custom fabrication", row[hDict['remarks']].value, flags=re.IGNORECASE))
                if ", " in temp:
                    run = run + temp[2:] + ", See Plans, Drawing #___ "
                else:
                    run = run + "See Plans, Drawing #___ "
            else:
                if(row[hDict['remarks']].value == None):
                    run = run + ("---")
                else:
                    run = run + (str(row[hDict['remarks']].value))

            #Utilities
            run = run + ("\nUtilities Req'd:\t")
            is_empty = True
            
#================= Imperial Utilities ================
                #Plumbing
            if not metric: 
                if(set(['hgt._','cw','d.w.']).issubset(hDict) and str(row[hDict['hgt._']].value)[0] == "-" and row[hDict['cw']].value == None):
                    run = run + (str(row[hDict['d.w.']].value) + " drain recessed " + str(row[hDict['hgt._']].value)[1:])
                    is_empty = False
            
                #Electrical
                if(set(['connection load', 'voltage', 'phase', 'comments__']).issubset(hDict) and row[hDict['connection load']].value != None):
                    a=str(row[hDict['connection load']].value).split("_x000D_\n")
                    v=str(row[hDict['voltage']].value).split("_x000D_\n")
                    ph=str(row[hDict['phase']].value).split("_x000D_\n")
                    c=None
                    if(row[hDict['comments__']].value != None):
                        c=str(row[hDict['comments__']].value).split("_x000D_\n")
                    for i in range(len(a)):
                        if not is_empty:
                            run = run + "; "
                        if "(" in str(a[i]):
                            run = run + str(a[i])[:3]+ " " 
                    
                        run = run + str(v[i]) + "/" + str(ph[i])+", "
                    
                        if "(" in str(a[i]):
                            run = run + str(a[i])[3:] 
                        else:
                            run = run + str(a[i])       
                
                        if (c != None):
                            if i<len(c):
                                run = run + " (" + str(c[i]) + ")"
                    is_empty = False
                
                    #Ventilation
                cfm = [] 
                unit = " CFM"
                if 'cfm' in hDict and row[hDict['cfm']].value != None:
                    if type(row[hDict['cfm']].value) == int:
                        cfm.append(str(row[hDict['cfm']].value)+ unit + " Exhaust")
                    else:
                        values = str("".join(str(row[hDict['cfm']].value).split("_x000D_"))).split()
                        if values.count(values[0]) == len(values):
                            cfm.append("(" + str(values.count(values[0])) + ")" + values[0] + unit +" Exhaust")
                        else:
                            for value in values:
                                cfm.append(value[:4] + unit +" Exhaust")
                if 'cfm_' in hDict and type(row[hDict['cfm_']].value) == str:
                    values = str("".join(str(row[hDict['cfm_']].value).split("_x000D_"))).split()
                    print(values)
                    if values.count(values[0]) == len(values):
                        cfm.append("(" + str(values.count(values[0])) + ")" + values[0] + unit +" Supply")
                    else:
                        for value in values:
                            cfm.append(value[:4] + unit +" Supply")
                
            
                if cfm:
                    if not is_empty:
                        run = run + "; "
                    run = run + ", ".join(cfm)
                    is_empty = False
                    
                tempList = []
            
                    #Plumbing (but more)   
                    #Water
                if 'cw' in hDict and row[hDict['cw']].value != None:
                    if("_x000D_" in str(row[hDict['cw']].value)):
                        temp = row[hDict['cw']].value.split("_x000D_")
                        if "\n" in "".join(temp):
                            temp = "".join(temp).split("\n")
                        if temp.count(temp[0]) == len(temp):
                            tempList.append("(" + str(temp.count(temp[0])) + ") " + temp[0] + " CW")
                        else:
                            tempList.append(", ".join(temp) + " CW")
                    else:
                        tempList.append(str(row[hDict['cw']].value) + " CW")

                if 'hw' in hDict and row[hDict['hw']].value != None:
                    tempList.append(str(row[hDict['hw']].value) + " HW")
                        
                    #Waste
                if 'waste' in hDict and row[hDict['waste']].value != None:
                    tempList.append(str(row[hDict['waste']].value) + " IW")

                if 'd.w.' in hDict and row[hDict['d.w.']].value != None and tempList:
                    tempList.append(str(row[hDict['d.w.']].value) + " DW")
                

                if not is_empty and tempList:
                    run = run + "; " 
                
                run = run + (", ".join(tempList))

                if tempList:
                    is_empty = False

                    #Gas
                if set(['size_','btu\'s','w.c.']).issubset(hDict) and row[hDict['size_']].value != None:
                    if not is_empty:
                        run = run + "; "
                    run = run + str(row[hDict['size_']].value) + " Gas @ " + str(row[hDict['btu\'s']].value) + " BTU; " + str(row[hDict['w.c.']].value) + " WC"
                    is_empty = False
           
                    #Chilled Water
                if set(['in size', 'out size']).issubset(hDict) and row[hDict['out size']].value != None:
                    if not is_empty:
                        run = run + "; "
                    else:
                        is_empty = False
                    run = run +str(row[hDict['in size']].value) + " Chilled Water Supply, " + str(row[hDict['out size']].value) + " Chilled Water Return"
                    
#================== METRIC UTILITIES ==========================
            else:
                if(set(['hgt. (mm)_','cw (mm)']).issubset(hDict) and row[hDict['hgt. (mm)_']].value != None and str(row[hDict['hgt. (mm)_']].value)[0] == "-" and row[hDict['cw (mm)']].value == None) and ("mm" not in str(row[hDict['hgt. (mm)_']].value)):
                    run = run + (str(row[hDict['hgt. (mm)_']].value) + "mm drain recessed " + str(row[hDict['cw (mm)']].value)[1:])
                    is_empty = False
        
                if(set(['connection load','voltage','phase']).issubset(hDict) and row[hDict['connection load']].value != None):
                    a=str(row[hDict['connection load']].value).split("_x000D_\n")
                    v=str(row[hDict['voltage']].value).split("_x000D_\n")
                    ph=str(row[hDict['phase']].value).split("_x000D_\n")
                    c=None
                    if('comments__' in hDict and row[hDict['comments__']].value != None):
                        c=str(row[hDict['comments__']].value).split("_x000D_\n")
                    for i in range(len(a)):
                        if not is_empty:
                            run = run + "; "
                        if "(" in str(a[i]):
                            run = run + str(a[i])[:3]+ " " 
                    
                        run = run + str(v[i]) + "/" + str(ph[i])+", "
                    
                        if "(" in str(a[i]):
                            run = run + str(a[i])[3:] 
                        else:
                            run = run + str(a[i])       
                
                        if (c != None):
                            if i<len(c):
                                run = run + " (" + str(c[i]) + ")"
                    is_empty = False
                unit = " M^3/H"
                cfm=[]
                if 'm^3/h' in hDict and row[hDict['m^3/h']].value != None:
                    if type(row[hDict['m^3/h']].value) == int:
                        cfm.append(str(row[hDict['m^3/h']].value)+ unit + " Exhaust")
                    else:
                        values = str("".join(str(row[hDict['m^3/h']].value).split("_x000D_"))).split()
                        if values.count(values[0]) == len(values):
                            cfm.append("(" + str(values.count(values[0])) + ")" + values[0] + unit +" Exhaust")
                        else:
                            for value in values:
                                cfm.append(value[:4] + unit +" Exhaust")
                if 'm^3/h_' in hDict and type(row[hDict['m^3/h_']].value) == str:
                    values = str("".join(str(row[hDict['m^3/h_']].value).split("_x000D_"))).split()
                    print(values)
                    if values.count(values[0]) == len(values):
                        cfm.append("(" + str(values.count(values[0])) + ")" + values[0] + unit +" Supply")
                    else:
                        for value in values:
                            cfm.append(value[:4] + unit +" Supply")

                tempList=[]
                if 'cw (mm)' in hDict and row[hDict['cw (mm)']].value != None:
                    if ("_x000D_" in str(row[hDict['cw (mm)']].value)):
                        temp = row[hDict['cw (mm)']].value.split("_x000D_")
                        if '\n' in "".join(temp):
                            temp = "".join(temp).split("\n")
                        if 'mm' not in temp[0]:
                            if temp.count(temp[0]) == len(temp):
                                tempList.append("(" + str(temp.count(temp[0])) + ") " + temp[0] + "mm CW")
                            else:
                                tempList.append(", ".join(temp) + "mm CW")
                        else:
                            if temp.count(temp[0]) == len(temp):
                                tempList.append("(" + str(temp.count(temp[0])) + ") " + temp[0] + " CW")
                            else:
                                tempList.append(", ".join(temp) + " CW")
                    else:
                        tempList.append(str(row[hDict['cw (mm)']].value) + " CW")

                if 'hw (mm)' in hDict and row[hDict['hw (mm)']].value != None:
                    if "mm" not in str(row[hDict['hw (mm)']].value):
                        tempList.append(str(row[hDict['hw (mm)']].value) + "mm HW")
                    else:
                        tempList.append(str(row[hDict['hw (mm)']].value) + " HW")

                if 'waste' in hDict and row[hDict['waste']].value != None and "mm" not in str(row[hDict['waste']].value):
                    tempList.append(str(row[hDict['waste']].value) + "mm IW")
                                    
                if 'd.w. (mm)' in hDict and row[hDict['d.w. (mm)']].value != None and tempList:
                    if "mm" not in str(row[hDict['d.w. (mm)']].value):
                        tempList.append(str(row[hDict['d.w. (mm)']].value) + "mm DW")
                    else:
                        tempList.append(str(row[hDict['d.w. (mm)']].value) + " DW")

                if not is_empty and tempList:
                    run = run + "; " 
                
                run = run + (", ".join(tempList))

                if tempList:
                    is_empty = False
        
                if set(['size (mm)_', 'mbar', 'kw']).issubset(hDict) and row[hDict['size (mm)_']].value != None:
                    if not is_empty:
                        run = run + "; "
                    run = run + str(row[hDict['size (mm)_']].value) + "mm LP Gas @ " + str(row[hDict['kw']].value) + " KW "
                    if row[hDict['mbar']].value != None:
                        run = run + "; " + str(row[hDict['mbar']].value) + " mbar"
                    is_empty = False

                if set(['in (mm)','out (mm)']).issubset(hDict) and row[hDict['out (mm)']].value != None:
                    if not is_empty:
                        run = run + "; "
                    else:
                        is_empty = False
                    if 'mm' not in str(row[hDict['in (mm)']].value):
                        run = run + str(row[hDict['in (mm)']].value) + "mm Chilled Water Supply, " + str(row[hDict['out (mm)']].value) + "mm Chilled Water Return"
                    else:
                        run = run + str(row[hDict['in (mm)']].value) + " Chilled Water Supply, " + str(row[hDict['out (mm)']].value) + " Chilled Water Return"

       
         
            if is_empty:
                run = run +("---")
        
            #Add Header to doc
            p.add_run(run)
                        
            ambiguousModels = ["custom", "custom design"] #Model No. that aren't specific to a model

            #Make/find specs for item
#SPECS BODY            
            #Existing Items Specs
            if row[hDict['remarks']].value != None and "EXIST" in str(row[hDict['remarks']].value):
                            
                p = doc.add_paragraph('', style = 'Spec_Header')
                remaining = (row[hDict['remarks']].value != None and "REMAIN" in str(row[hDict['remarks']].value))           
                name = ""
                if row[hDict['description']].value != None:
                    name = row[hDict['description']].value.lower()
                specText = ""
                if(remaining):
                    specText = specText + "Remain in place existing unit as follows:\n"
                else:
                    specText = specText + "Relocate existing unit as follows:\n"
                specList = []
                if remaining:
                    specList.append("Existing unit is located in existing kitchen; unit should be thoroughly cleaned and remaing where shown on plan")
                else:
                    specList.append("Existing unit is located in existing kitchen; unit should be thoroughly cleaned and relocated where shown on plan")
                    specList.append("Schedule time with Owner for relocating unit")
                if "shelv" in name.lower():
                    specList.append("Replace shelves where corrosion spots appear; clean, sand, polish and repaint if necessary")
                elif "trash" not in name.lower() and "bin" not in name.lower():
                    specList.append("Repair where corrosion spots appear; clean, sand, polish and repain if necessary")
                i = 1
                for spec in specList:
                    specText = specText + str(i) + ".\t" + spec + "\n"
                    i = i+1
                p.add_run(specText)
                rr = p.add_run(str(i) + ".\tVerify all existing utility requirements and conditions\n"+ str(i+1) + ".\tThoroughly clean and sanitize unit\n")
                rr.font.color.rgb = d.shared.RGBColor(0xFF,0x00,0x00)
                p.add_run(str(i+2) + ".\tMust meet all applicable federal, state, and local laws, rules, regulations and codes")

                
            #If excel sheet is not provided, search through DB for a match
            elif excelFilepath == "":
                con = None 
                dbPath = resource_path('specsDB.db')
                try:
                    con=sqlite3.connect("file:"+dbPath+"?mode=rw", uri=True)
                except sqlite3.OperationalError:
                    con=sqlite3.connect(dbPath)
                    cur = con.cursor()
                    cur.execute("CREATE TABLE IF NOT EXISTS item (desc, manu, model, doc)")
                    cur.execute("CREATE TABLE IF NOT EXISTS spec (doc, text, modTime)")
                    con.commit()
                cur = con.cursor()
                specData = []
                #Manually fill field for custom fab
                if type(row[hDict['remarks']].value) == str and "CUSTOM FABRICATION" in row[hDict['remarks']].value.upper():
                    specData = [row[hDict['description']].value, "Custom Fabrication", ""]
                #Fill fields with Excel values
                else:
                    specData = [row[hDict['description']].value, row[hDict['manuf.']].value, str(row[hDict['model']].value).replace('/', '-').replace('|','-')]
                
                matches = []
                if specData[1] == "Custom Fabrication":
                    matches = cur.execute("SELECT * FROM item WHERE desc='" + str(specData[0]).replace("'","''").replace('"','""') + "' COLLATE NOCASE AND manu = 'Custom Fabrication' COLLATE NOCASE").fetchall()
                else:
                    matches = cur.execute("SELECT * FROM item WHERE model='" + str(specData[2]).replace("'","''").replace('"','""') + "'").fetchall()
                
                if matches:
                    specDoc = ""
                    #If there were multiple matches, offer user to choose the best match
                    
                    if len(matches) > 1:
                        
                        try:
                            #Try using dictionary to make decision
                            specDoc = specDict[str(specData[2])]
                            
                        except:
                            #Have user make a decision and save it to dictionary
                            top=tk.Toplevel(root)
                            top.geometry("800x400")
                            top.title("Select a Source Spec")
                            tk.Label(top,text="Multiple Specs Match the Following Item.\nDescription: "+ str(specData[0]) +"\nManufacturer: "+ str(specData[1]) +"\nModel No.: "+ str(specData[2]) +"\n\n\nPlease Choose Best Match").pack()
                            columns=["Manufacturer","Model","Word Doc","Path"]
                            treeview = ttk.Treeview(top, selectmode = 'browse', columns =columns)
                            treeview.pack()
                            matchedString = tk.StringVar(root, "")
                            select = tk.Button(top, text="Select", command=lambda:ChooseSpec(matchedString, top, treeview))
                            select.pack()
                            
                            treeview.heading("#0", text="Description", command=lambda: treeview_sort_column(treeview, '#0', True))
                            treeview.heading("Manufacturer", text="Manufacturer",command=lambda: treeview_sort_column(treeview, columns[0], True))
                            treeview.heading("Model", text="Model",command=lambda: treeview_sort_column(treeview, columns[1], True))
                            treeview.heading("Word Doc", text="Word Doc",command=lambda: treeview_sort_column(treeview, columns[2], True))
                            treeview.column('#4', stretch='no', minwidth=0, width=0, anchor=tk.E)
                            treeview.tag_bind('tag?', "<Double-1>", lambda event, msg = msgLabel: openFile(event, msg))
                            for entry in matches:        
                                t= treeview.insert('', tk.END, text =str(entry[0]), values = (str(entry[1]), str(entry[2]), str(entry[3]).split('/')[len(str(entry[3]).split('/'))-1],str(entry[3])), tags=("tag?",)) 
                            root.wait_window(top)
                            specDoc = matchedString.get()
                            specDict[str(specData[2])] = specDoc
                    else:
                        specDoc = matches[0][3]
                    
                    copySpecs(specDoc, doc.add_paragraph('', style = 'Spec_Header'), False, cur)
                else:
                    #Check partial matches
                    if (specData[1] != "Custom Fabrication" and str(specData[2]).lower() not in ambiguousModels):
                        matches = cur.execute("SELECT desc, model, doc FROM item WHERE model LIKE '%" + str(specData[2]).replace("'","''").replace('"','""') + "%' AND manu LIKE '%" + str(specData[1]).replace("'","''").replace('"','""') + "%'").fetchall()
                    if not matches: 
                        matches = cur.execute("SELECT desc, model, doc FROM item WHERE desc LIKE '%" + str(specData[0]).replace("'","''").replace('"','""')+ "%' AND manu LIKE '%"+ str(specData[1]).replace("'","''").replace('"','""') +"%'").fetchall()
                    if matches:
                        def closestMatch(s1, sList, i):
                            minDist = max(len(sList[0][i]), len(s1))
                            closest = sList[0]
                            for s in sList:
                                temp = edit_distance(s1, s[i])
                                if temp < minDist:
                                    minDist = temp
                                    closest = s
                            return closest

                        bestMatch = matches[0]
                        if len(matches) > 1:
                            if specData[2] == 'Stainless Steel':
                                bestMatch = closestMatch(specData[0], matches, 0)
                            else:
                                bestMatch = closestMatch(specData[2], matches, 1)
                
              
                        copySpecs(bestMatch[2], doc.add_paragraph('', style = 'Spec_Header'), True, cur)

            #If Spec Ref Sheet has been provided
            else:
                con = None 
                dbPath = resource_path('specsDB.db')
                try:
                    con=sqlite3.connect("file:"+dbPath+"?mode=rw", uri=True)
                except sqlite3.OperationalError:
                    con=sqlite3.connect(dbPath)
                    cur = con.cursor()
                    cur.execute("CREATE TABLE IF NOT EXISTS item (desc, manu, model, doc)")
                    cur.execute("CREATE TABLE IF NOT EXISTS spec (doc, text, modTime)")
                    con.commit()
                cur = con.cursor()
                #if specs exist, copy and paste    
                if (row[hDict['model']].value != None and str(row[hDict['model']].value).lower() in specRefs[2] and specRefs[4][specRefs[2].index(str(row[hDict['model']].value).lower())]
                      and str(row[hDict['model']].value).lower() not in ambiguousModels) and specRefs[3][specRefs[2].index(str(row[hDict['model']].value).lower())] != "":
                    
                    if -1 == copySpecs(specRefs[3][specRefs[2].index(str(row[hDict['model']].value).lower())], doc.add_paragraph('', style = 'Spec_Header'), False, cur) and specRefs[3][specRefs[2].index(str(row[hDict['model']].value).lower())] not in broken:
                        broken.append(specRefs[3][specRefs[2].index(str(row[hDict['model']].value).lower())])
                    
                # If Manufacturer and Desc. match, copy specs
                elif (row[hDict['description']].value != None and str(row[hDict['description']].value).lower() in specRefs[0]):
                    manuf = ""
                    #Check if specs exist for matching Manufacturer and Desc.
                    if(row[hDict['manuf.']].value != None and str(row[hDict['manuf.']].value).lower() in specRefs[1]):
                        manuf = str(row[hDict['manuf.']].value).lower()
                    #Check if specs for custom fabrication item exists
                    elif(row[hDict['remarks']].value != None and "custom fabrication" in str(row[hDict['remarks']].value).lower()):
                        manuf = "custom fabrication"
                    #If neither, skip
                    else:
                        continue
                    #Find specs which match for given item
                    for index in [i for i,e in enumerate(specRefs[1]) if e == manuf]:
                        if index in [j for j,s in enumerate(specRefs[0]) if s == str(row[hDict['description']].value).lower()]: 
                            if(specRefs[3][index] == ""):
                                continue
                            #Copy and paste from associated doc and highlighting from spec ref sheet
                            if -1 == copySpecs(specRefs[3][index], doc.add_paragraph('', style = 'Spec_Header'),not specRefs[4][index],cur) and specRefs[3][index] not in broken:
                                broken.append(specRefs[3][index])
                            break                           

    wb.close()
    try:
        doc.save(outputFilepath+"/Specs.docx")
    except:
        msgLabel.config(text="Error: Cannot save Specs.docx while file is open")
        return
    if broken:
        msgLabel.config(text= "The following doc file(s) could not be found: \n" + '\n'.join(broken))
    elif msgLabel.cget("text") == 'Working...':
        msgLabel.config(text="Successfully Created Specs Document")
    print(str(time.time()-start_time) + " secs")
#GUI

#Select Input File
def getFilepath(inputLabel):
    global inputFilepath
    inputFilepath = filedialog.askopenfilename(filetypes = (("Microsoft Excel Worksheet", "*.xlsx"),))
    inputLabel.config(text= "The input file is: " + inputFilepath)

#Select Output Location
def getOutputFolder(outputLabel):
    global outputFilepath 
    outputFilepath = filedialog.askdirectory()
    outputLabel.config(text= "The output folder is: " + outputFilepath)

#Select Excel File
def getExcelFile(xlLabel):
    global excelFilepath 
    excelFilepath = filedialog.askopenfilename(filetypes = (("Microsoft Excel Worksheet", "*.xlsx"),))
    xlLabel.config(text= "The ref sheet location is: " + excelFilepath)

#Search DB for anything that partially matches user input
def Search(tv, text, cur, var, msg):
    fields = ["","","",""]
    match var.get():
        case "Model":
            fields[2] = text
        case "Description":
            fields[0] = text
        case "Manufacturer":
            fields[1] = text
        case _:
            print("How did you do that?????")
    if fields[0] or fields[1] or fields[2] or fields[3]:
        try:
            found = db.FindEntry(fields, cur)
            msg.config(text = "Search successful")
        except:
            
            msg.config(text = "Avoid using special characters")
            return
        tv.delete(*tv.get_children())
        for item in found:
            tv.insert("", tk.END, text = item[0], values = (item[1], item[2], item[3].split('/')[len(item[3].split('/'))-1], item[3]),tags = ("tag?",))
    else:
        tv.delete(*tv.get_children())
        entries = cur.execute("SELECT * FROM item").fetchall()
        for entry in entries:
            tv.insert("", tk.END, text = entry[0], values = (entry[1], entry[2], entry[3].split('/')[len(entry[3].split('/'))-1], entry[3]),tags=("tag?",))

def ModifyEntry(con, tv, root, msg):
    cur = con.cursor()
    item = None
    try:
        item = tv.selection()[0]
    except:
        msg.config(text="Please select an item")
        return
    
    itemInfo = tv.item(item)

    #popup with info for changes
    top=tk.Toplevel(root)
    top.geometry("500x400")
    top.title("Edit Spec")

    docPath=tk.StringVar(root,"")
    docPath.set(itemInfo['values'][3])

    cornerBuffer = tk.Label(top)
    cornerBuffer.grid(row=0, column=0, padx=35, pady=10)
    centerBuffer = tk.Label(top)
    centerBuffer.grid(row=3, column=2, padx=40, pady=10)
    
    descText = tk.Text(top, height=1, width=15)
    descText.insert(tk.INSERT, itemInfo['text'])
    descText.grid(row=2,column=1)
    manuText = tk.Text(top, height=1, width=15)
    manuText.insert(tk.INSERT, itemInfo['values'][0])
    manuText.grid(row=2,column=3)
    modelText = tk.Text(top, height=1, width=15)
    modelText.grid(row=5,column=1)
    modelText.insert(tk.INSERT, itemInfo['values'][1])

    descLabel = tk.Label(top, text="Enter Description:")
    descLabel.grid(row=1, column=1)
    manuLabel = tk.Label(top, text="Enter Manufacturer:")
    manuLabel.grid(row=1, column=3)
    modelLabel = tk.Label(top, text="Enter Model:")
    modelLabel.grid(row=4, column=1)
    docLabel = tk.Label(top, text="Choose Doc File:")
    docLabel.grid(row=4, column=3)
    docPathLabel = tk.Label(top, text="Spec Doc is located at: " + docPath.get(), wraplength=200)
    docPathLabel.grid(row=6, column=1, columnspan=3, pady=10)

    msgLabel = tk.Label(top,text="",wraplength=400)
    msgLabel.grid(row=8,column=1,columnspan=3,padx=10,pady=10)

    def chooseDoc():
        docPath.set(filedialog.askopenfilename(filetypes = (("Microsoft Word Document", "*.docx"),)))
        top.lift()
        docPathLabel.configure(text="Spec Doc is located at: " + docPath.get())

    docButton = tk.Button(top, text="Select Doc", command=chooseDoc)
    docButton.grid(row=5,column=3)
    ch = tk.StringVar(root, "")
    res=tk.StringVar(root, "0")

    def submit():
        changes = [descText.get('1.0', 'end-1c'), manuText.get('1.0', 'end-1c'), modelText.get('1.0', 'end-1c'), docPath.get()]
        badChars = ['\\','/','|','?',':','*','<','>','"',"'",',']
        ch.set(",".join(changes))
        
        if any(char in ' '.join(changes[:3]) for char in badChars):
            msgLabel.config(text="Warning: Cannot use the following characters: " + ','.join(badChars))
            return
        if any(char in changes[3] for char in ('"',"'",',')):
            msgLabel.config(text="Warning: Cannot open Doc file due to special characters in the name.")
            return
        res.set(str(db.ModifyEntry([itemInfo["text"], itemInfo["values"][0], itemInfo["values"][1], itemInfo["values"][3]], changes, cur)))    
        top.destroy()

    submitButton = tk.Button(top, text="Submit", command=submit)
    submitButton.grid(row=7,column=2, pady = 20)

    root.wait_window(top)
    if int(res.get()) <1:
        msg.config(text="Couldn\'t update entry")
        return
    elif int(res.get()) > 1:
        msg.config(text="Modified multiple entries")
    else:
        msg.config(text="Successfully Modified Entry")
    changes = ch.get().split(",")
    if len(changes) < 4:
        return
    con.commit()
    
#Edit Treeview
    vals = ['','','','']
    
    if changes[0]:
        desc = changes[0]
    else:
        desc = itemInfo['text']
        
    if changes[1]:
        vals[0] = changes[1]
    else:
        vals[0] = itemInfo['values'][0]

    if changes[2]:
        vals[1] = changes[2]
    else:
        vals[1] = itemInfo['values'][1]

    if changes[3]:
        vals[2] = changes[3].split('/')[len(changes[3].split('/'))-1]
        vals[3] = changes[3]
    else:
        vals[2] = itemInfo['values'][2]
        vals[3] = itemInfo['values'][3]
    tv.item(item, text=desc, values=vals)
    
def AddEntry(con, tv, msg):
    cur = con.cursor()
    missing = []
    for f in filedialog.askopenfilenames():
        if any(char in f for char in ('"',"'",',')):
            missing.append(f)
            continue
        splitfile = f.split('/')
        
        splitfile = splitfile[len(splitfile)-1].split('_')
        if len(splitfile) < 3:
            missing.append(f)
            continue
        
        if db.addEntry([splitfile[0].strip(), splitfile[1].strip(), splitfile[2].split('.docx')[0].strip(), f], cur):
            tv.insert('', tk.END, text =str(splitfile[0].strip()), values = (splitfile[1].strip(), splitfile[2].split('.docx')[0].strip(), '_'.join(splitfile), f), tags=("tag?",))
        
       
    if missing:
        msg.config(text='Couldn\'t add specs for the following (doesn\'t follow naming convention):\n '+'\n'.join(missing))
    else:
        msg.config(text='Successfully added entry(s)')
    con.commit()
        
    
#Temporary Placeholder Function
def Nothing():
    pass

#Remove entry from treeview and DB
def DeleteEntry(con, tv, msg):
    cur = con.cursor()
    try:
        item = tv.selection()[0]
    except:
        msg.config(text="No item selected")
        return
    itemInfo = tv.item(item)
    tv.delete(item)
    
    res = db.DeleteEntry(itemInfo['values'][3], cur)
   
    if res[0] >1 or res[1]>1:
        msg.config(text = "Deleted Multiple Entries with doc: " + itemInfo['values'][3])
    elif res[0] and res[1]:  
        msg.config(text = "Successfully Deleted Entry")
    else:
        msg.config(text="Couldn't find entry: " + itemInfo['values'][3])
    con.commit()
    
    
#Update all out-of-date specs with changes in Doc file
def UpdateSpecs(con,msg):
    cur = con.cursor()
    missing = db.UpdateSpecs(cur)
   
    if missing:
        msg.config(text="Couldn\'t find the following files: " + "\n".join(missing))
    else:
        msg.config(text="Successfully updated specs")
    con.commit()
    
#Open file from treeview
def openFile(event, msg):
    tree=event.widget
    item = tree.item(tree.focus())
    try:
        os.startfile(item['values'][3])
    except:
        msg.config(text="This file doesn't exist here anymore...")

#Sort treeview based on selected column
def treeview_sort_column(tv, col, reverse):
    l=[]
    if col == '#0':
        l = [(tv.item(k)["text"],k) for k in tv.get_children('')]
    else:
        l = [(tv.set(k,col),k) for k in tv.get_children('')]
    l.sort(key = lambda t: t[0],reverse=reverse)
    for index, (val, k) in enumerate(l):
        tv.move(k,'',index)
    tv.heading(col, command=lambda: treeview_sort_column(tv,col, not reverse))
    
#Create window for choosing functionality
def selectionWindow():
    for widget in root.winfo_children():
        widget.destroy()
    root.geometry("800x500")

    #Help Menu
    menubar = tk.Menu(root)
    helpMenu = tk.Menu(menubar, tearoff=0)
    helpMenu.add_command(label = "Help", command = lambda:os.startfile(resource_path('Help.html')))
    helpMenu.add_command(label = "Examples", command = lambda:os.startfile(filedialog.askopenfilename(initialdir=resource_path("Example Files"))))
    menubar.add_cascade(label="Help",menu=helpMenu)
    root.config(menu=menubar)
        

    frame = tk.Frame(root)
    frame.pack()
    
    findSpecButton = tk.Button(frame, text="Manage DB", command=DBWindow, height=2, width=15, bg= '#afafaf')
    findSpecButton.pack(padx=80,pady=150,side=tk.LEFT)

    writeSpecButton = tk.Button(frame, text="Write Specs", command=wsWindow, height=2, width=15, bg= '#afafaf')
    writeSpecButton.pack(padx=80,pady=100,side=tk.RIGHT)

#Window for displaying and interacting with DB
def DBWindow():
    for widget in root.winfo_children():
        widget.destroy()
        
    root.geometry("1100x500")
   

    menubar = tk.Menu(root)
    helpMenu = tk.Menu(menubar, tearoff=0)
    helpMenu.add_command(label = "Help", command = lambda:os.startfile(resource_path('Help.html')))
    helpMenu.add_command(label = "Examples", command = lambda:os.startfile(filedialog.askopenfilename(initialdir=resource_path("./Example Files"))))
    menubar.add_cascade(label="Help",menu=helpMenu)
    root.config(menu=menubar)

    backButton = tk.Button(root, text="<-", command=selectionWindow, bg = '#dadada')
    backButton.grid(row=0, column=0, sticky = 'W', ipadx=15,ipady=5)

    messageBox = tk.Label(root, text = "")
    messageBox.grid(row=6,column=2, columnspan=4 ,pady=10,padx=10)
    
    columns = ("Manufacturer", "Model", "Word Doc", "Path")
    treeview = ttk.Treeview(root, selectmode = 'browse', columns =columns)
    treeview.heading("#0", text="Description", command=lambda: treeview_sort_column(treeview, '#0', True))
    treeview.heading("Manufacturer", text="Manufacturer",command=lambda: treeview_sort_column(treeview, columns[0], True))
    treeview.heading("Model", text="Model",command=lambda: treeview_sort_column(treeview, columns[1], True))
    treeview.heading("Word Doc", text="Word Doc",command=lambda: treeview_sort_column(treeview, columns[2], True))
    treeview.column('#4', stretch='no', minwidth=0, width=0, anchor=tk.E)
    treeview.tag_bind('tag?', "<Double-1>", lambda event, msg=messageBox:openFile(event, msg))

    con = None 
    dbPath = resource_path("specsDB.db")
    try:
        con=sqlite3.connect("file:"+dbPath+"?mode=rw", uri=True)
    except sqlite3.OperationalError:
        con=sqlite3.connect(dbPath)
        cur = con.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS item (desc, manu, model, doc)")
        cur.execute("CREATE TABLE IF NOT EXISTS spec (doc, text, modTime)")
        con.commit()
    cur = con.cursor()
    entries = cur.execute("SELECT * FROM item").fetchall()

    for entry in entries:        
        t= treeview.insert('', tk.END, text =str(entry[0]), values = (str(entry[1]), str(entry[2]), entry[3].split("/")[len(entry[3].split("/"))-1], entry[3]), tags=("tag?",))

    scrollBar = ttk.Scrollbar(root, orient="vertical", command = treeview.yview)
    scrollBar.grid(row=2,column=6,rowspan=4, sticky='ns')
    treeview.grid(row=2,column=2, rowspan=4, columnspan=4)
    treeview.configure(yscrollcommand =scrollBar.set)

    

    add = tk.Button(root, text= "Add Entry", command=lambda:AddEntry(con,treeview,messageBox))
    add.grid(row=2, column=1, padx = 20)

    modify = tk.Button(root, text = "Modify Entry", command=lambda:ModifyEntry(con, treeview, root, messageBox))
    modify.grid(row=3, column=1, padx = 20)

    delete = tk.Button(root, text = "Delete Entry", command=lambda:DeleteEntry(con, treeview, messageBox))
    delete.grid(row=4, column=1, padx = 20)

    update = tk.Button(root, text = "Update Entries", command=lambda:UpdateSpecs(con, messageBox))
    update.grid(row=5, column=1, padx = 20)
    
    searchFrame = tk.Frame(root)
    searchFrame.grid(row=1, column=2, columnspan=4, pady=20)

    search = tk.Label(searchFrame, text= "Search for ")
    search.grid(row=0, column=0)

    T = tk.Text(searchFrame, height=1.2, width=15)
    T.grid(row=0, column=1)

    search2 = tk.Label(searchFrame, text=" in ")
    search2.grid(row=0, column=2)

    options = ["Description", "Manufacturer", "Model"]
    var = tk.StringVar(root)
    var.set("Model")

    drop = tk.OptionMenu(searchFrame, var, *options)
    drop.grid(row=0, column=3)
    
    submit = tk.Button(searchFrame, text="Search", command=lambda:Search(treeview, T.get("1.0", 'end-1c'), cur, var, messageBox))
    submit.grid(row=0, column=4, columnspan=2, padx = 20, pady=20)
    def SearchEvent(event, tree, text,cur,var,msg):
        Search(tree,text,cur,var,msg)
        return "break"
    T.bind('<Return>', lambda event:SearchEvent(event, treeview, T.get("1.0",'end-1c'), cur, var, messageBox))

#Create Window for Writing Specs Function
def wsWindow():
    root.geometry("800x600")
    for widget in root.winfo_children():
        widget.destroy()

        
    menubar = tk.Menu(root)
    helpMenu = tk.Menu(menubar, tearoff=0)
    helpMenu.add_command(label = "Help", command = lambda:os.startfile(resource_path('Help.html')))
    helpMenu.add_command(label = "Examples", command = lambda:os.startfile(filedialog.askopenfilename(initialdir=resource_path("./Example Files"))))
    menubar.add_cascade(label="Help",menu=helpMenu)
    root.config(menu=menubar)
    backButton = tk.Button(root, text="<-", command=selectionWindow, bg = '#dadada')
    backButton.grid(row=0, column=0, sticky = 'W', ipadx=15,ipady=5)

    padLabel = tk.Label(root, text="")
    padLabel.grid(row=10,column=1, padx=80, pady=30)

    messageLabel = tk.Label(root, text="")
    messageLabel.grid(row=9, column = 2, columnspan = 3)

    infoLabel = tk.Label(root, text="Write Specs Document")
    infoLabel.grid(row=1, column=3, sticky= 'N', pady=30)

    inputLabel = tk.Label(root, wraplength = 250, text="The input file is: " + inputFilepath)
    inputLabel.grid(row=5, column=2, columnspan =3)

    inputButton = tk.Button(root, text="Select Input File", command=lambda:getFilepath(inputLabel))
    inputButton.grid(row=3, column=2, padx=10, pady=30, rowspan=2)

    outputLabel = tk.Label(root, wraplength = 250, text="The output location is: " + outputFilepath)
    outputLabel.grid(row=6, column=2, columnspan =3)

    outputButton = tk.Button(root, text="Select Output Folder", command=lambda:getOutputFolder(outputLabel))
    outputButton.grid(row=3, column=4, padx=10, pady=10, rowspan=2)

    xlLabel = tk.Label(root, wraplength = 250, text="The ref sheet location is: " + excelFilepath)
    xlLabel.grid(row=7, column=2, columnspan =3)
    
    xlButton = tk.Button(root, text="Select Ref Sheet\n(Optional)", command=lambda:getExcelFile(xlLabel))
    xlButton.grid(row=4, column=3, padx=10, pady=10)

    refButton = tk.Button(root, text="Make Ref Sheet\n(Optional)", command=lambda:findSpecs(messageLabel))
    refButton.grid(row=3, column=3, padx=10, pady=10)
  
    submitButton = tk.Button(root, text="Create Specs", command=lambda:writeSpecs(messageLabel))
    submitButton.grid(row=8, column = 3, pady = 20)

#Open the selection window
selectionWindow()

root.mainloop()
