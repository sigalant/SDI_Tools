import docx

masterDoc = docx.Document("AutoText Master.docx")
newDoc = docx.Document()

fullText = []

para = newDoc.add_paragraph('')

finished = False

for p in masterDoc.paragraphs:
    if "ITEM" in p.text:
        fileName = p.text.split('\t')[len(p.text.split('\t'))-1].strip().replace('/','-').replace('|','-')
    elif "Manu" in p.text:
        fileName = fileName + "_" + p.text.split('\t')[len(p.text.split('\t'))-1].strip().replace('/','-').replace('|','-')
    elif "Model No." in p.text:
        fileName = fileName + "_" + p.text.split('\t')[len(p.text.split('\t'))-1].strip().replace('/','-').replace('|','-')
    p_runs = []
    add_runs = False
    beginning = True
    for run in p.runs:
        if run.font.color.rgb != docx.shared.RGBColor(0x00,0x00,0x00):

            if fullText:
                para.add_run('\n'.join(fullText) + "\n")
                fullText = []
            if beginning:
                beginning = False
            if p_runs:
                para.add_run(''.join(p_runs))
            p_runs = []
            addRuns = True
            redRun = para.add_run(run.text)
            redRun.font.color.rgb = run.font.color.rgb
        else:
            p_runs.append(run.text)
    if run._element.br_lst:
        finished = True    
    if addRuns:
        para.add_run(''.join(p_runs))
        addRuns = False
    else:
        fullText.append(p.text)
    if not fullText:
        para.add_run('\n')
    p_runs = []
    if finished:
        try:
            newDoc.save("./Word Files/" + fileName + ".docx")
        except:
            print(fileName)
        finished = False
        fileName = ""
        newDoc = docx.Document()
        para = newDoc.add_paragraph('')